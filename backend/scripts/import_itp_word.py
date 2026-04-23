"""
从原版 ITP Word 文档导入 itp_definitions 和 rfi_groundfields。
提取每个 .docx 中「6. ITP TABLE」与「7. REVISION CHANGE DETAILS」之间的表格，
解析后写入数据库（与 Excel 导入脚本一致的数据结构）。

用法（项目根目录下，使用虚拟环境 myenv）:
  python -m backend.scripts.import_itp_word "D:\\Inspections\\ITP\\word"
  python -m backend.scripts.import_itp_word "D:\\Inspections\\ITP\\word" --preview
  python -m backend.scripts.import_itp_word "D:\\Inspections\\ITP\\word" --dry-run
  python -m backend.scripts.import_itp_word "D:\\Inspections\\ITP\\word" --file "GCC-CC7-PM-00000-QC-PLN-00301_04土建工程ITP.docx"

  不传 --file 时：处理目录下全部 .docx（全量导入）；传 --file 时只处理该文件。
  --preview      仅预览每个文档中提取到的表格（表头+前几行）
  --dry-run      不提交数据库，仅打印将写入的数据
  --no-skip-itp  已存在的 ITP 也更新 itp_name/version/status（现与 rfi_groundfields 一致，重导即更新）
  rfi_groundfields 采用更新/插入/删除策略（按 document_number + level + itp_id 匹配），不整表删除。
"""
import sys
import os
import re
import argparse
from pathlib import Path
from typing import Optional, Any, List, Tuple, Iterator

project_root = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "backend"))
os.chdir(project_root)

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from sqlalchemy import delete
from sqlalchemy.orm import Session
from app.database import SessionLocal
from app.models.report import ITPDefinition, RFIGroundField


# 原版 Word 表头多为英文/俄文，列名匹配（包含即匹配）
COLUMN_ALIASES = {
    "itp_id": ["No.", "No", "Item No", "№", "序号", "编号"],
    "section_name": ["Section", "大节", "章节"],
    "workdescription": ["Activity description", "Work description", "工作描述", "工作内容", "Description"],
    "applicable_documents": ["Applicable documents", "Applicable document", "适用文件", "规范", "标准", "程序"],
    "acceptance_criteria": ["Acceptance criteria", "Acceptance criterion", "验收准则", "验收标准"],
    "quality_control_form": ["Quality control form", "Quality control", "verifying document", "质量控制", "记录表", "主控文件", "Master document"],
    "master_document": ["Master document", "主文件"],
    "involvement": ["Involvement", "参与", "Participant", "P/R/S"],
    "involvement_subcon": ["Subcon", "Subсontractor", "Sub-contractor", "分包商", "Subcontractor", "Субподрядчик"],
    "involvement_contractor": ["Contractor", "承包商", "总包", "Подрядчик"],
    "involvement_customer": ["Customer", "Owner", "业主", "业主/CQC", "Заказчик"],
    "involvement_aqc": ["AQC"],  # 移除了 CQC 以避免误匹配 Customer / CQC
}


def _iter_block_items(document: Document) -> Iterator[Any]:
    """按文档顺序遍历段落和表格（Paragraph / Table）。"""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:tbl"):
            yield Table(child, document)
        elif child.tag == qn("w:p"):
            yield Paragraph(child, document)


def _get_num_id_and_lvl(p: Paragraph) -> Tuple[Optional[int], Optional[int]]:
    """从段落或其样式（含继承样式）中提取 numId 和 ilvl。"""
    num_id_xpath = ".//w:numPr/w:numId/@w:val"
    ilvl_xpath = ".//w:numPr/w:ilvl/@w:val"
    try:
        # 1. 检查段落直接属性
        num_ids = p._element.xpath(num_id_xpath)
        if num_ids:
            ilvls = p._element.xpath(ilvl_xpath)
            return int(num_ids[0]), (int(ilvls[0]) if ilvls else 0)
        # 2. 检查样式及其父样式
        curr_style = p.style
        while curr_style:
            num_ids = curr_style.element.xpath(num_id_xpath)
            if num_ids:
                ilvls = curr_style.element.xpath(ilvl_xpath)
                return int(num_ids[0]), (int(ilvls[0]) if ilvls else 0)
            if not hasattr(curr_style, 'base_style') or not curr_style.base_style:
                break
            curr_style = curr_style.base_style
    except:
        pass
    return None, None


def _paragraph_text(p: Paragraph, numbering_map: Optional[dict] = None) -> str:
    """仅供 extract_itp_tables 用于解析标题/标记，不处理单元格内部编号重置。"""
    t = (p.text or "").strip()
    if numbering_map and p._element in numbering_map:
        num_prefix = numbering_map[p._element]
        prefix_pattern = re.escape(num_prefix).replace(r"\.", r"\.?")
        if not re.match(r"^" + prefix_pattern, t) and not re.match(r"^\d+[\.\)\-\s]", t):
            t = f"{num_prefix} {t}"
    return t


def _cell_text(cell, numbering_map: Optional[dict] = None, is_id_column: bool = False) -> str:
    """提取单元格文本。"""
    if not cell.paragraphs:
        return (cell.text or "").strip()
    
    parts = []
    # 单元格内统一计数器
    cell_counters = [0] * 10
    
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        num_id, ilvl = _get_num_id_and_lvl(p)
        
        # 智能嗅探：如果文本开头已经带了数字编号（手写），同步计数器
        # 匹配 "1.", "1)", "1 " 等格式
        manual_match = re.match(r"^(\d+)[\.\)\s]", t)
        if manual_match and not is_id_column:
            try:
                cell_counters[ilvl] = int(manual_match.group(1))
            except: pass

        if num_id is not None:
            num_prefix = None
            if is_id_column:
                if numbering_map and p._element in numbering_map:
                    num_prefix = numbering_map[p._element]
            else:
                # 只有当文本本身没有编号开头时，才增加自动编号前缀
                # 哪怕手写和自动混合，这里也能强制连续
                if not manual_match:
                    cell_counters[ilvl] += 1
                    # 当当前层级增加时，重置所有子层级
                    for i in range(ilvl + 1, 10):
                        cell_counters[i] = 0
                    
                    active = [str(cell_counters[i]) for i in range(ilvl + 1) if cell_counters[i] > 0]
                    if active:
                        num_prefix = ".".join(active)
                        if ilvl == 0: num_prefix += "."
            
            if num_prefix:
                # 如果文本已经有编号，就不再重复加
                if not re.match(r"^\d+[\.\)\-\s]", t):
                    if is_id_column and not t:
                        t = num_prefix
                    else:
                        t = f"{num_prefix} {t}"
        parts.append(t)
    return "\n".join(parts).strip()


def _build_numbering_map(doc: Document) -> dict:
    """建立全局编号映射，主要用于维持 No. 列跨单元格的连贯性。"""
    numbering_map = {}
    counters = {}
    visited_paragraphs = set()

    def process_paragraph(p):
        if p._element in visited_paragraphs:
            return
        visited_paragraphs.add(p._element)
        num_id, ilvl = _get_num_id_and_lvl(p)
        if num_id is None:
            return
        try:
            if num_id not in counters:
                counters[num_id] = [0] * 10
            counters[num_id][ilvl] += 1
            for i in range(ilvl + 1, 10):
                counters[num_id][i] = 0
            active = [str(counters[num_id][i]) for i in range(ilvl + 1) if counters[num_id][i] > 0]
            if not active: return
            num_text = ".".join(active)
            if ilvl == 0: num_text += "."
            numbering_map[p._element] = num_text
        except: pass

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            process_paragraph(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_paragraph(p)
    return numbering_map


def _table_contains_marker(tbl: Table, pattern: re.Pattern, numbering_map: Optional[dict] = None) -> bool:
    """检查表格任意单元格是否包含匹配 pattern 的文本。"""
    for row in tbl.rows:
        for cell in row.cells:
            if pattern.search(_cell_text(cell, numbering_map)):
                return True
    return False


def _extract_itp_name_from_doc(doc: Document, numbering_map: Optional[dict] = None) -> Optional[str]:
    """从正文、页眉、页脚取 ITP 名称：含 Inspection and test plan / План инспекций 的段落。"""
    def _check_text(text: str) -> Optional[str]:
        if not text or len(text) < 10:
            return None
        if "Inspection and test plan" in text or "План инспекций" in text or "Inspection and Test Plan" in text:
            return text[:500].strip()
        return None

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            out = _check_text(_paragraph_text(block, numbering_map))
            if out:
                return out
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    out = _check_text(_cell_text(cell, numbering_map))
                    if out:
                        return out
    for section in doc.sections:
        for story in (section.header, section.footer):
            if story is None:
                continue
            for p in story.paragraphs:
                out = _check_text((p.text or "").strip())
                if out:
                    return out
    return None


def _extract_version_from_revision_table(tbl: Table, numbering_map: Optional[dict] = None) -> Optional[str]:
    """从 7. REVISION CHANGE DETAILS 表中取 Rev No. / Изм.№ 列的第一个数据行（如 04）。"""
    rows = [[_cell_text(c, numbering_map) for c in tr.cells] for tr in tbl.rows]
    if len(rows) < 2:
        return None
    header = rows[0]
    rev_col = None
    for i, cell in enumerate(header):
        c = (cell or "").strip().upper()
        if "REV NO" in c or "ИЗМ" in c or "REVISION" in c:
            rev_col = i
            break
    if rev_col is None:
        return None
    for r in rows[1:]:
        if rev_col < len(r):
            val = (r[rev_col] or "").strip()
            if val and re.match(r"^\d+$", val):
                return val[:50]
    return None


def _parse_stem_for_version_and_name(stem: str) -> Tuple[str, Optional[str], Optional[str]]:
    """
    从文件名 stem 解析 document_number、version、itp_name。
    例如 GCC-CC7-PM-00000-QC-PLN-00301_04土建工程ITP -> document_number, "04", "土建工程ITP"
    """
    if "_" not in stem:
        return stem, None, stem
    document_number = stem.split("_")[0].strip()
    rest = stem.split("_", 1)[1].strip()
    if not rest:
        return document_number, None, document_number
    ver_match = re.match(r"^(\d+)", rest)
    version = ver_match.group(1)[:50] if ver_match else None
    itp_name = re.sub(r"^\d+", "", rest).strip() or rest
    return document_number, version, itp_name


def extract_itp_tables(doc_path: Path) -> Tuple[Optional[str], Optional[str], Optional[str], List[List[List[str]]]]:
    """
    打开 Word，定位「6. ITP TABLE」与「7. REVISION CHANGE DETAILS」之间的表格。
    返回 (document_number, itp_name, version, tables)。
    """
    doc = Document(str(doc_path))
    numbering_map = _build_numbering_map(doc)
    stem = doc_path.stem
    document_number, version_from_file, itp_name_from_file = _parse_stem_for_version_and_name(stem)
    itp_name = itp_name_from_file or document_number
    version: Optional[str] = version_from_file

    in_section = False
    tables_in_section: List[Table] = []
    start_marker = re.compile(r"6\.\s*ITP\s+TABLE", re.I)
    end_marker = re.compile(r"7\.\s*REVISION\s+CHANGE\s+DETAIL", re.I)

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = _paragraph_text(block, numbering_map)
            if start_marker.search(text):
                in_section = True
                continue
            if in_section and end_marker.search(text):
                break
            continue
        if isinstance(block, Table):
            if not in_section:
                if _table_contains_marker(block, start_marker, numbering_map):
                    in_section = True
                    if not _table_contains_marker(block, end_marker, numbering_map):
                        tables_in_section.append(block)
                continue
            tables_in_section.append(block)
            if _table_contains_marker(block, end_marker, numbering_map):
                if version is None:
                    version = _extract_version_from_revision_table(block, numbering_map)
                break
    
    if not tables_in_section:
        blocks = list(_iter_block_items(doc))
        start_idx = None
        end_idx = None
        for i, block in enumerate(blocks):
            if isinstance(block, Paragraph):
                t = start_marker.search(_paragraph_text(block, numbering_map))
                if t and start_idx is None: start_idx = i
                t = end_marker.search(_paragraph_text(block, numbering_map))
                if t:
                    end_idx = i
                    if start_idx is not None: break
            elif isinstance(block, Table):
                if _table_contains_marker(block, start_marker, numbering_map) and start_idx is None:
                    start_idx = i
                if _table_contains_marker(block, end_marker, numbering_map):
                    end_idx = i
                    if version is None: version = _extract_version_from_revision_table(block, numbering_map)
                    if start_idx is not None: break
        if start_idx is not None and end_idx is not None:
            for j in range(start_idx, end_idx):
                if isinstance(blocks[j], Table):
                    tables_in_section.append(blocks[j])

    result_tables = []
    for tbl in tables_in_section:
        # 先简单确定 No. 列索引
        itp_id_col_idx = None
        if tbl.rows:
            header_cells = [str(c.text or "").strip().upper() for c in tbl.rows[0].cells]
            for i, c in enumerate(header_cells):
                if any(a.upper() in c for a in COLUMN_ALIASES["itp_id"]):
                    itp_id_col_idx = i
                    break
        
        rows = []
        for tr in tbl.rows:
            row_vals = []
            prev_cell_el = None
            for i, cell in enumerate(tr.cells):
                is_id = (i == itp_id_col_idx)
                # 处理水平合并：如果当前单元格与前一个是同一个物理单元格（水平合并），则在该行后续列中视为空
                if prev_cell_el is not None and cell._element == prev_cell_el:
                    row_vals.append("")
                else:
                    row_vals.append(_cell_text(cell, numbering_map, is_id_column=is_id))
                prev_cell_el = cell._element
            rows.append(row_vals)
        if rows:
            result_tables.append(rows)
            
    result_tables = [t for t in result_tables if _is_itp_data_table(t)]
    return document_number, itp_name, version, result_tables


def _is_itp_data_table(rows: List[List[str]]) -> bool:
    """判断是否为 ITP 数据表（有 No.、工作描述、适用文件等列），排除 PURPOSE、REVISION、ATTACHMENTS 表。"""
    if not rows:
        return False
    first_row_text = " ".join(str(c or "") for c in rows[0]).upper()
    if "PURPOSE AND SCOPE" in first_row_text or "REVISION CHANGE" in first_row_text or "ATTACHMENTS" in first_row_text:
        return False
    # 前几行中需出现 No./№ 与 Work description/Activity description 与 Applicable
    for row in rows[:4]:
        row_text = " ".join(str(c or "") for c in row)
        row_upper = row_text.upper()
        has_no = "NO." in row_upper or "№" in row_text or ("NO" in row_upper and "№" in row_text)
        has_work = "WORK DESCRIPTION" in row_upper or "ACTIVITY DESCRIPTION" in row_upper or "ОПИСАНИЕ РАБОТ" in row_upper
        has_applicable = "APPLICABLE" in row_upper or "ПРИМЕНИМЫЕ" in row_upper
        if has_no and has_work and has_applicable:
            return True
    return False


def _find_header_row(rows: List[List[str]], max_header_rows: int = 4) -> int:
    """找到表头行索引：第一个同时包含 No./№ 与 Work/Activity description 的行。"""
    for i in range(min(max_header_rows, len(rows))):
        row = rows[i]
        row_upper = " ".join(str(c or "") for c in row).upper()
        has_no = "NO." in row_upper or "№" in row_upper
        has_work = "WORK DESCRIPTION" in row_upper or "ACTIVITY DESCRIPTION" in row_upper
        if has_no and has_work:
            return i
    return 0


def _find_column_index(header_row: List[str], field: str) -> Optional[int]:
    """在表头行中找匹配 COLUMN_ALIASES 的列索引。使用更严格的单词边界匹配。"""
    aliases = COLUMN_ALIASES.get(field, [field])
    for i, cell in enumerate(header_row):
        cell_clean = (cell or "").strip().lower()
        if field == "itp_id" and ("description" in cell_clean or "описание" in cell_clean):
            continue
        for a in aliases:
            if not a: continue
            al = a.lower()
            # 只有当别名是单元格内容的完整部分（单词边界）时才匹配
            pattern = r'\b' + re.escape(al) + r'(\b|\.)'
            if re.search(pattern, cell_clean) or al == cell_clean:
                return i
    return None


# 根据全量文档分析得到的精确表头短语
HEADER_ROW_PHRASES = (
    "No.", "№", "No. №", "Item No",
    "Work description", "Activity description", "Описание работ",
    "Applicable documents", "Applicable documents (Specifications, codes & procedures)",
    "Acceptance criteria", "Quality control form", "Quality control form, master document",
    "Quality control form, verifying document", "Involvement", "Involvement Участие",
    "Involvement / Участие", "Participant", "P/R/S",
    "Subcontractor", "Contractor", "Customer", "Customer / CQC", "AQC",
    "Применимые документы", "Критерии приёмки", "Форма контроля", "Участие",
    "Субподрядчик", "Подрядчик", "Заказчик",
    "CONTENTS", "СОДЕРЖАНИЕ", "PURPOSE AND SCOPE", "НАЗНАЧЕНИЕ И ОБЛАСТЬ ПРИМЕНЕНИЯ",
    "REVISION CHANGE DETAILS", "REVISION CHANGE DETAILEDS", "ЛИСТ ИЗМЕНЕНИЙ",
    "ATTACHMENTS", "ПРИЛОЖЕНИЯ"
)


def _cell_looks_like_header(cell_val: Any) -> bool:
    """单元格内容是否为表头短语。必须是精确匹配（不计大小写和换行），严禁‘包含匹配’。"""
    s = str(cell_val or "").strip()
    if not s or len(s) > 100:
        return False
    
    # 统一格式：去掉换行和多余空格
    s_norm = re.sub(r'\s+', ' ', s).upper()
    
    for phrase in HEADER_ROW_PHRASES:
        p_norm = re.sub(r'\s+', ' ', phrase).upper()
        # 要求全字匹配
        if s_norm == p_norm:
            return True
    return False


def _looks_like_section_header(s: str) -> bool:
    """判断字符串是否像大节标题（如 1. xxx、2) xxx、3- xxx），用于 No. 列未合并时从 Work description 等列识别。"""
    if not s or not str(s).strip():
        return False
    s = str(s).strip()
    # 数字 + 可选 . ) - + 至少一个空格或紧跟字母（含俄文等），且不能是 1.1 这种
    if re.match(r"^\d+\.\d+", s):
        return False
    return bool(re.match(r"^\d+(\.|[\)\-])?(\s+|[A-Za-z\u0400-\u04ff])", s))


def _first_section_header_from_row(row: List[str], max_cols: int = 2) -> Optional[str]:
    """从行内前 max_cols 个单元格中取第一个「像大节标题」的文本。限制扫描前2列，避免误判适用文件列。"""
    for j in range(min(max_cols, len(row))):
        cell = row[j]
        if cell is None:
            continue
        s = str(cell or "").strip()
        if _looks_like_section_header(s):
            return s
    return None


def _normalize_level_from_no(no_val: Any, work_or_section_val: Any = None, row_cells: Optional[List[str]] = None) -> int:
    """根据 No. / 工作描述列判断 level：
    - 仅数字/数字+点（如 1 或 1.）视为 level 2（大节）
    - 1.1 / 1.2 这类视为 level 3（检查项）
    - No. 为空时，若工作描述/大节列或以 row_cells 任一侧以「1. xxx」「1 xxx」「1) xxx」「1- xxx」等开头也视为 level 2。
    支持从 Work description 开始合并单元格的表格（No. 列不合并时用整行前几列扫描）。
    """
    s_no = str(no_val or "").strip()
    if s_no:
        # 1.1 / 1.2 / 10.3.4 等 -> level 3
        if re.match(r"^\d+\.\d+", s_no):
            return 3
        # 1 / 1. / 1 . / 1 .   -> level 2
        if re.match(r"^\d+\.?\s*$", s_no) or re.match(r"^\d+\s*\.", s_no):
            return 2

    # No. 为空或未命中时，看工作描述/大节列或整行前几列是否有大节标题
    s_work = str(work_or_section_val or "").strip()
    if _looks_like_section_header(s_work):
        return 2
    if row_cells:
        first_section = _first_section_header_from_row(row_cells)
        if first_section:
            return 2

    return 3


def _extract_section_number_from_text(s: Optional[str]) -> Optional[str]:
    """从大节标题文本中提取序号，如 '2. Material...' -> '2.'，'3) xxx' -> '3.'。"""
    if not s or not str(s).strip():
        return None
    s = str(s).strip()
    m = re.match(r"^(\d+)(\.|[\)\-])?", s)
    return m.group(1) + "." if m else None


def _normalize_itp_id_for_level(level: int, no_val: Any, section_header_text: Optional[str] = None) -> Optional[str]:
    """Level 2 的 itp_id 规范化（保留 4. 或 4.1 这种形式）。Level 3 的 itp_id 规范化（去重复、去尾点）。"""
    s = str(no_val or "").strip()
    if level == 2:
        if s:
            # 提取开头的数字部分（支持 1 或 1.1 或 1.1.1）
            m = re.match(r"^(\d+(?:\.\d+)*)", s)
            if m:
                val = m.group(1)
                # 如果只是单个数字（如 4），补全为 4. 以符合旧版习惯
                if "." not in val:
                    val += "."
                return val
            return None
        if section_header_text:
            return _extract_section_number_from_text(section_header_text)
        return None
    if not s:
        return None
    
    # Level 3 规范化
    # 1. 移除重复前缀，如 "1.1 1.1" -> "1.1"
    parts = s.split()
    if len(parts) >= 2 and parts[0].rstrip('.') == parts[1].rstrip('.'):
        s = parts[1] # 保留后者，通常手动部分更准确
    # 2. 移除尾部点，统一格式便于匹配（1.3. 和 1.3 视为一致）
    s = re.sub(r"[\.\s]+$", "", s)
    return _str(s, 255)


def _to_json_array(val: Any) -> Optional[list]:
    if val is None or (isinstance(val, str) and not val.strip()):
        return None
    s = str(val).strip()
    parts = re.split(r"[\n;；]", s)
    out = [p.strip() for p in parts if p.strip()]
    return out if out else None


def _str(val: Any, max_len: int = 1048576) -> Optional[str]:
    if val is None or (isinstance(val, str) and not val.strip()):
        return None
    s = str(val).strip()[:max_len]
    return s or None


def _parse_involvement(val: Any) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]:
    """把一列参与方拆成 4 个 P/R/S 值（按顺序：分包商、承包商、业主、AQC）。"""
    if val is None or (isinstance(val, str) and not val.strip()):
        return None, None, None, None
    parts = re.split(r"[\s,]+", str(val).strip())
    parts = [_str(p, 10) for p in parts if p and str(p).strip()][:4]
    while len(parts) < 4:
        parts.append(None)
    return parts[0], parts[1], parts[2], parts[3]


def _has_cyrillic(s: Any) -> bool:
    """判断字符串是否包含俄文字符。"""
    if not s:
        return False
    return bool(re.search(r'[\u0400-\u04ff]', str(s)))


def table_rows_to_ground_rows(
    document_number: str,
    itp_name: str,
    all_tables_rows: List[List[List[str]]],
) -> List[dict]:
    """将一个文档内提取到的所有表格行转为 ground_rows。支持跨表/跨页合并。"""
    ground_rows = []
    level2_order = 0
    level3_order = 0
    
    # 按照顺序收集所有的有效数据行
    all_data_rows = []
    for table_rows in all_tables_rows:
        if len(table_rows) < 2:
            continue
        header_row_idx = _find_header_row(table_rows)
        header = [str(c or "").strip() for c in table_rows[header_row_idx]]
        
        col_map = {}
        for field in COLUMN_ALIASES:
            idx = _find_column_index(header, field)
            if idx is not None:
                col_map[field] = idx
        
        # 补全四列索引
        invol_col = col_map.get("involvement")
        sub_row_idx = header_row_idx + 1
        if invol_col is not None and sub_row_idx < len(table_rows):
            sub_header = [str(c or "").strip() for c in table_rows[sub_row_idx]]
            for f in ("involvement_subcon", "involvement_contractor", "involvement_customer", "involvement_aqc"):
                if col_map.get(f) is None:
                    idx = _find_column_index(sub_header, f)
                    if idx is not None:
                        col_map[f] = idx

        def get_cell(row, col_idx):
            if col_idx is None or col_idx >= len(row): return None
            return row[col_idx]

        for i in range(header_row_idx + 1, len(table_rows)):
            row = table_rows[i]
            
            # --- 终极防御逻辑：跳过跨页重复的表头行 ---
            matches = 0
            # 检查关键列特征
            for field in ["itp_id", "workdescription", "applicable_documents", "acceptance_criteria"]:
                c_idx = col_map.get(field)
                if c_idx is not None and c_idx < len(row):
                    if _cell_looks_like_header(row[c_idx]):
                        matches += 1
            
            # 如果两个及以上关键单元格完全匹配表头短语，这行就是“阴险”的重复表头
            if matches >= 2:
                continue
            
            # 特殊情况：序号列直接就是明显的表头字符
            no_val = str(get_cell(row, col_map.get("itp_id")) or "").strip()
            if no_val in ("No.", "№", "No"):
                continue
            # --------------------------------------
                
            all_data_rows.append((row, col_map))

    # 按顺序处理所有收集到的行
    pending_l3: Optional[dict] = None
    last_itp_id_level3: Optional[str] = None

    for row, col_map in all_data_rows:
        itp_id_col = col_map.get("itp_id")
        section_col = col_map.get("section_name")
        work_col = col_map.get("workdescription")
        app_doc_col = col_map.get("applicable_documents")
        acc_cri_col = col_map.get("acceptance_criteria")
        qc_col = col_map.get("quality_control_form")
        master_col = col_map.get("master_document")
        invol_col = col_map.get("involvement")
        subcon_col = col_map.get("involvement_subcon")
        contractor_col = col_map.get("involvement_contractor")
        customer_col = col_map.get("involvement_customer")
        aqc_col = col_map.get("involvement_aqc")

        def get_cell_val(c_idx):
            if c_idx is None or c_idx >= len(row): return None
            return row[c_idx]

        no_val = get_cell_val(itp_id_col)
        work_or_section = get_cell_val(work_col) or get_cell_val(section_col)
        first_section_in_row = _first_section_header_from_row(row)
        work_or_section = work_or_section or first_section_in_row
        
        section_val = _str(get_cell_val(section_col))
        work_val = _str(get_cell_val(work_col))
        app_doc_val = _to_json_array(get_cell_val(app_doc_col))
        acc_cri_val = _to_json_array(get_cell_val(acc_cri_col))
        qc_val = _str(get_cell_val(qc_col))
        master_val = _str(get_cell_val(master_col))
        
        # Involvement 处理
        def get_inv():
            if subcon_col is not None:
                subcon_cell = get_cell_val(subcon_col)
                parsed = _parse_involvement(subcon_cell)
                if all(parsed): return parsed
                return (_str(subcon_cell, 10), _str(get_cell_val(contractor_col), 10), _str(get_cell_val(customer_col), 10), _str(get_cell_val(aqc_col), 10))
            if invol_col is not None:
                single = get_cell_val(invol_col)
                parsed = _parse_involvement(single)
                if all(parsed): return parsed
                vals = [_str(get_cell_val(invol_col + i), 10) for i in range(4)]
                while len(vals) < 4: vals.append(None)
                return tuple(vals)
            return None, None, None, None

        subcon_val, contractor_val, customer_val, aqc_val = get_inv()

        # 初始判定 Level
        level = _normalize_level_from_no(no_val, work_or_section, row_cells=row)
        
        # 核心逻辑：如果原本判定为信息行(Level 3)，但所有关键数据列都是空的，则识别为分组(Level 2)
        if level == 3:
            # 数据列包括：适用文件、验收标准、质量控制、主文档、以及四个参与方列
            is_data_empty = not (app_doc_val or acc_cri_val or qc_val or master_val or subcon_val or contractor_val or customer_val or aqc_val)
            if is_data_empty and (work_val or section_val):
                level = 2

        itp_id_val = _normalize_itp_id_for_level(level, no_val, section_header_text=work_or_section if level == 2 else None)
        
        if level == 3 and itp_id_val is None and last_itp_id_level3 is not None:
            itp_id_val = last_itp_id_level3
            
        if level == 2:
            if pending_l3:
                level3_order += 1
                pending_l3["sort_order"] = level3_order
                ground_rows.append(pending_l3)
                pending_l3 = None
            last_itp_id_level3 = None
            level2_order += 1
            level3_order = 0
            section_itp_id = itp_id_val if itp_id_val else (str(level2_order) + ".")
            ground_rows.append({
                "document_number": document_number, "level": 2, "parent_id": None, "itp_id": section_itp_id,
                "description": itp_name, "section_name": work_val or section_val, "sort_order": level2_order,
                "workdescription_eng": None, "workdescription_rus": None, "workdescription_chn": None,
                "applicable_documents_eng": None, "applicable_documents_rus": None, "applicable_documents_chn": None,
                "acceptance_criteria_eng": None, "acceptance_criteria_rus": None, "acceptance_criteria_chn": None,
                "quality_control_form_master_document_eng": None, "quality_control_form_master_document_rus": None,
                "quality_control_form_master_document_chn": None, "involvement_subcon": None,
                "involvement_contractor": None, "involvement_customer": None, "involvement_aqc": None,
            })
        else:
            # --- 终极语言判定逻辑：精准区分英俄行 ---
            # 核心：只根据“工作描述”列来判断。避免因为“适用文件”列引用了俄语标准（如 СП, ГОСТ）而导致整行被误判为俄语。
            is_rus = _has_cyrillic(work_val)
            
            # 如果工作描述为空（极少数跨页补丁行），则回退到检查章节名或其他内容
            if not work_val:
                is_rus = _has_cyrillic(section_val) or _has_cyrillic(get_cell_val(app_doc_col))
            # --------------------------------------

            # 改进的合并逻辑：只要 ID 匹配就合并，无论语言顺序，也无论是否跨页
            if pending_l3 and pending_l3.get("itp_id") == itp_id_val:
                if is_rus:
                    # 合并/补充俄文内容
                    pending_l3["workdescription_rus"] = _merge_text_fields(pending_l3.get("workdescription_rus"), work_val)
                    pending_l3["applicable_documents_rus"] = _merge_list_fields(pending_l3.get("applicable_documents_rus"), app_doc_val)
                    pending_l3["acceptance_criteria_rus"] = _merge_list_fields(pending_l3.get("acceptance_criteria_rus"), acc_cri_val)
                    pending_l3["quality_control_form_master_document_rus"] = _merge_text_fields(pending_l3.get("quality_control_form_master_document_rus"), qc_val or master_val)
                else:
                    # 合并/补充英文内容
                    pending_l3["workdescription_eng"] = _merge_text_fields(pending_l3.get("workdescription_eng"), work_val)
                    pending_l3["applicable_documents_eng"] = _merge_list_fields(pending_l3.get("applicable_documents_eng"), app_doc_val)
                    pending_l3["acceptance_criteria_eng"] = _merge_list_fields(pending_l3.get("acceptance_criteria_eng"), acc_cri_val)
                    pending_l3["quality_control_form_master_document_eng"] = _merge_text_fields(pending_l3.get("quality_control_form_master_document_eng"), qc_val or master_val)
                
                # 参与方只要有新值就更新
                if subcon_val: pending_l3["involvement_subcon"] = subcon_val
                if contractor_val: pending_l3["involvement_contractor"] = contractor_val
                if customer_val: pending_l3["involvement_customer"] = customer_val
                if aqc_val: pending_l3["involvement_aqc"] = aqc_val
            else:
                # ID 不匹配，先提交上一个
                if pending_l3:
                    level3_order += 1
                    pending_l3["sort_order"] = level3_order
                    ground_rows.append(pending_l3)
                
                # 开始新的 pending
                pending_l3 = {
                    "document_number": document_number, "level": 3, "parent_id": None, "itp_id": itp_id_val,
                    "description": itp_name, "section_name": section_val, "sort_order": level3_order,
                    "workdescription_eng": work_val if not is_rus else None,
                    "workdescription_rus": work_val if is_rus else None,
                    "applicable_documents_eng": app_doc_val if not is_rus else None,
                    "applicable_documents_rus": app_doc_val if is_rus else None,
                    "acceptance_criteria_eng": acc_cri_val if not is_rus else None,
                    "acceptance_criteria_rus": acc_cri_val if is_rus else None,
                    "quality_control_form_master_document_eng": (qc_val or master_val) if not is_rus else None,
                    "quality_control_form_master_document_rus": (qc_val or master_val) if is_rus else None,
                    "involvement_subcon": subcon_val, "involvement_contractor": contractor_val,
                    "involvement_customer": customer_val, "involvement_aqc": aqc_val,
                    "workdescription_chn": None, "applicable_documents_chn": None, "acceptance_criteria_chn": None, "quality_control_form_master_document_chn": None,
                }
            if itp_id_val: last_itp_id_level3 = itp_id_val

    if pending_l3:
        level3_order += 1
        pending_l3["sort_order"] = level3_order
        ground_rows.append(pending_l3)
    return ground_rows


def _ground_row_key(r: dict) -> Tuple[str, int, Optional[str]]:
    """用于跨页合并：同一 (document_number, level, itp_id/section_name) 的多行合并为一条。"""
    doc = r.get("document_number")
    level = r.get("level", 3)
    itp_id = r.get("itp_id")
    if level == 2:
        sec = (r.get("section_name") or "").strip()
        return (doc, 2, itp_id or sec or None)
    return (doc, level, itp_id)


def _merge_list_fields(a: Optional[list], b: Optional[list]) -> Optional[list]:
    """合并两个列表（用于 applicable_documents / acceptance_criteria），去重保留顺序。"""
    if not a and not b:
        return None
    seen = set()
    out = []
    for x in (a or []) + (b or []):
        if x is None:
            continue
        s = str(x).strip()
        if s and s not in seen:
            seen.add(s)
            out.append(x)
    return out if out else None


def _merge_text_fields(*vals: Any) -> Optional[str]:
    """多段文本用换行拼接，去空。"""
    parts = []
    for v in vals:
        if v is None or (isinstance(v, str) and not str(v).strip()):
            continue
        s = str(v).strip()
        if s:
            parts.append(s)
    return "\n".join(parts) if parts else None


def _merge_ground_rows(rows: List[dict]) -> dict:
    """将同一 key 的多行（跨页重复）合并为一条：列表类字段扩展合并，文本类换行拼接，标量取首个非空。"""
    if len(rows) == 1:
        return rows[0].copy()
    first = rows[0]
    merged = dict(first)
    for r in rows[1:]:
        for lang in ("_eng", "_rus", "_chn"):
            ak = "applicable_documents" + lang
            merged[ak] = _merge_list_fields(merged.get(ak), r.get(ak))
            ck = "acceptance_criteria" + lang
            merged[ck] = _merge_list_fields(merged.get(ck), r.get(ck))
        merged["workdescription_eng"] = _merge_text_fields(merged.get("workdescription_eng"), r.get("workdescription_eng"))
        merged["workdescription_rus"] = _merge_text_fields(merged.get("workdescription_rus"), r.get("workdescription_rus"))
        merged["workdescription_chn"] = _merge_text_fields(merged.get("workdescription_chn"), r.get("workdescription_chn"))
        for lang in ("_eng", "_rus", "_chn"):
            qk = "quality_control_form_master_document" + lang
            merged[qk] = _merge_text_fields(merged.get(qk), r.get(qk))
        for inv in ("involvement_subcon", "involvement_contractor", "involvement_customer", "involvement_aqc"):
            if not _str(merged.get(inv), 10) and _str(r.get(inv), 10):
                merged[inv] = r.get(inv)
    return merged


def merge_cross_page_duplicates(all_ground_rows: List[dict]) -> List[dict]:
    """按 (document_number, level, itp_id/section_name) 分组，同 key 多行（跨页）合并为一条，保留首次出现顺序。"""
    # 按文档，再按首次出现顺序收集 key -> [rows]
    doc_order: List[str] = []
    doc_keys_order: dict = {}  # doc -> list of keys in first-occurrence order
    groups: dict = {}  # (doc, level, key_val) -> list of rows
    for r in all_ground_rows:
        doc, level, key_val = _ground_row_key(r)
        k = (doc, level, key_val)
        if doc not in doc_order:
            doc_order.append(doc)
        if doc not in doc_keys_order:
            doc_keys_order[doc] = []
        if k not in groups:
            doc_keys_order[doc].append(k)
        groups.setdefault(k, []).append(r)
    out = []
    for doc in doc_order:
        for k in doc_keys_order[doc]:
            out.append(_merge_ground_rows(groups[k]))
    return out


def preview_word_dir(dir_path: Path, single_file: Optional[str] = None) -> None:
    """预览目录下每个 docx 中 6/7 之间的表格。"""
    files = list(dir_path.glob("*.docx"))
    files = [f for f in files if not f.name.startswith("~$")]
    if single_file:
        # 支持模糊匹配，处理编码或特殊字符问题
        files = [f for f in files if (f.name == single_file or single_file in f.name)]
    if not files:
        print(f"未找到 .docx 文件: {dir_path}")
        return
    for f in sorted(files):
        print(f"\n========== {f.name} ==========")
        try:
            doc_no, itp_name, version, tables = extract_itp_tables(f)
            print(f"  document_number: {doc_no}, version: {version}, itp_name: {(itp_name or '')[:60]}...")
            print(f"  表格数: {len(tables)}")
            for ti, tbl in enumerate(tables):
                print(f"  --- Table {ti + 1} ({len(tbl)} 行) ---")
                for ri, row in enumerate(tbl[:6]):
                    print(f"    Row {ri}: {row[:12]}")
                if len(tbl) > 6:
                    print(f"    ... 共 {len(tbl)} 行")
        except Exception as e:
            print(f"  错误: {e}")
            import traceback
            traceback.print_exc()


def import_from_word(
    dir_path: Path,
    single_file: Optional[str] = None,
    dry_run: bool = False,
    skip_existing_itp: bool = True,
    verbose: bool = False,
) -> None:
    """遍历目录下 docx，提取 6/7 间表格，写入 itp_definitions 和 rfi_groundfields。"""
    files = list(dir_path.glob("*.docx"))
    files = [f for f in files if not f.name.startswith("~$")]
    if single_file:
        # 支持模糊匹配，处理编码或特殊字符问题
        files = [f for f in files if (f.name == single_file or single_file in f.name)]
    if not files:
        print(f"未找到 .docx 文件: {dir_path}")
        return

    all_itp_rows = []
    all_ground_rows = []
    for f in sorted(files):
        try:
            doc_no, itp_name, version, tables = extract_itp_tables(f)
        except Exception as e:
            print(f"跳过 {f.name}: {e}")
            continue
        all_itp_rows.append({
            "document_number": doc_no,
            "itp_name": itp_name,
            "version": version,
            "status": "active",
        })
        n_ground_file = 0
        rows = table_rows_to_ground_rows(doc_no, itp_name, tables)
        all_ground_rows.extend(rows)
        n_ground_file += len(rows)
        if verbose:
            print(f"  {f.name}: {len(tables)} 张表 -> {n_ground_file} 条 ground 行")

    # 诊断：若最终写入 0 条，先看解析阶段是否得到行
    if not all_ground_rows and not dry_run:
        print("诊断: 解析阶段未得到任何 ground 行（表格被过滤或 table_rows_to_ground_rows 返回空）。可用 --verbose 查看每文件表数/行数。")
    elif verbose:
        print("解析得到 ground 行数（合并前）: %d" % len(all_ground_rows))

    # 跨页合并：同一文档内相同 (level, itp_id/section_name) 的多行合并为一条，再写入库
    n_before = len(all_ground_rows)
    all_ground_rows = merge_cross_page_duplicates(all_ground_rows)
    n_after = len(all_ground_rows)
    if n_before > n_after and not dry_run:
        print("跨页合并: 原始 %d 行 -> 合并后 %d 行（同编号多行已合并为一条）" % (n_before, n_after))

    if dry_run:
        print("\n[DRY RUN] 将写入 itp_definitions:", len(all_itp_rows), "条")
        for r in all_itp_rows[:10]:
            print(" ", r)
        print("\n[DRY RUN] 将写入 rfi_groundfields:", len(all_ground_rows), "条")
        for r in all_ground_rows[:5]:
            print(" ", r)
        if len(all_ground_rows) > 5:
            print(" ... 等共", len(all_ground_rows), "条")
        return

    db = SessionLocal()
    try:
        for r in all_itp_rows:
            existing = db.query(ITPDefinition).filter(ITPDefinition.document_number == r["document_number"]).first()
            if existing:
                # 重导时始终用本次解析结果更新 itp_name/version/status，与 skip_existing_itp 无关
                existing.itp_name = r["itp_name"]
                existing.version = r.get("version")
                existing.status = r.get("status", "active")
            else:
                db.add(ITPDefinition(
                    document_number=r["document_number"],
                    itp_name=r["itp_name"],
                    version=r.get("version"),
                    status=r.get("status", "active"),
                ))
        db.commit()
        print("已写入/更新 itp_definitions:", len(all_itp_rows), "条")

        # 按 (document_number, level, itp_id) 更新/插入/删除，不再整表删除后插入
        def _ground_attrs(r: dict, parent_id: Optional[int], sort_order: int) -> dict:
            # 按 DB 列长截断，避免 Data too long：section_name/description 255，quality_control_* 500
            return {
                "document_number": r["document_number"],
                "level": r["level"],
                "parent_id": parent_id,
                "itp_id": _str(r.get("itp_id"), 255),
                "description": r.get("description"),
                "section_name": r.get("section_name"),
                "workdescription_eng": r.get("workdescription_eng"),
                "workdescription_rus": r.get("workdescription_rus"),
                "workdescription_chn": r.get("workdescription_chn"),
                "applicable_documents_eng": r.get("applicable_documents_eng"),
                "applicable_documents_rus": r.get("applicable_documents_rus"),
                "applicable_documents_chn": r.get("applicable_documents_chn"),
                "acceptance_criteria_eng": r.get("acceptance_criteria_eng"),
                "acceptance_criteria_rus": r.get("acceptance_criteria_rus"),
                "acceptance_criteria_chn": r.get("acceptance_criteria_chn"),
                "quality_control_form_master_document_eng": r.get("quality_control_form_master_document_eng"),
                "quality_control_form_master_document_rus": r.get("quality_control_form_master_document_rus"),
                "quality_control_form_master_document_chn": r.get("quality_control_form_master_document_chn"),
                "involvement_subcon": r.get("involvement_subcon"),
                "involvement_contractor": r.get("involvement_contractor"),
                "involvement_customer": r.get("involvement_customer"),
                "involvement_aqc": r.get("involvement_aqc"),
                "sort_order": sort_order,
            }

        def _ground_key_from_row(r: dict) -> Tuple[int, Optional[str]]:
            """用于匹配/去重的 key。
            - level=2：优先用 itp_id；若 itp_id 为空，则退回用 section_name 文本，避免所有 itp_id=NULL 的大节被误认为同一条。
            - level=3：仍按 (level, itp_id) 匹配。
            """
            level = r.get("level")
            itp_id = r.get("itp_id")
            if level == 2:
                sec = (r.get("section_name") or "").strip()
                return 2, (itp_id or sec or None)
            return level, itp_id

        def _ground_key_from_rec(rec: RFIGroundField) -> Tuple[int, Optional[str]]:
            level = rec.level
            itp_id = rec.itp_id
            if level == 2:
                sec = (rec.section_name or "").strip()
                return 2, (itp_id or sec or None)
            return level, itp_id

        updated_count = 0
        inserted_count = 0
        deleted_count = 0
        # (document_number, level, key, reason)
        updated_details: List[Tuple[str, int, Tuple[int, Optional[str]], str]] = []
        for doc in sorted({r["document_number"] for r in all_ground_rows}):
            doc_rows = [r for r in all_ground_rows if r["document_number"] == doc]
            incoming_key_set = set(_ground_key_from_row(r) for r in doc_rows)
            existing_list = db.query(RFIGroundField).filter(RFIGroundField.document_number == doc).all()
            existing_map = {_ground_key_from_rec(r): r for r in existing_list}
            # 本文档本批中已插入过的 key：用于区分「同文档重复编号」与「与库内已有记录匹配」
            inserted_keys_this_doc: set = set()

            level2_id = None
            level2_order = 0
            level3_order = 0
            for r in doc_rows:
                key = _ground_key_from_row(r)
                if r["level"] == 2:
                    level2_order += 1
                    level3_order = 0
                    parent_id = None
                    sort_order = level2_order
                else:
                    level3_order += 1
                    parent_id = level2_id
                    sort_order = level3_order

                attrs = _ground_attrs(r, parent_id, sort_order)
                rec = existing_map.get(key)
                if rec:
                    for k, v in attrs.items():
                        setattr(rec, k, v)
                    updated_count += 1
                    if key in inserted_keys_this_doc:
                        reason = "同文档内重复编号：本表/本批中该 itp_id 出现多次，后行覆盖前行（请检查 Word 是否有多处相同编号或跨页重复）"
                    else:
                        reason = "与库内已有记录匹配（重导或历史数据），已覆盖该行"
                    updated_details.append((doc, r["level"], key, reason))
                    if r["level"] == 2:
                        level2_id = rec.id
                else:
                    rec = RFIGroundField(**attrs)
                    db.add(rec)
                    db.flush()
                    inserted_count += 1
                    inserted_keys_this_doc.add(key)
                    existing_map[key] = rec  # 同批次内重复 (level, itp_id) 时后续走更新，避免重复插入
                    if r["level"] == 2:
                        level2_id = rec.id

            # 将本循环中的更新/插入先刷入库，避免删除时 DB 仍认为子节点指向待删父节点
            db.flush()

            to_delete = [existing_map[key] for key in existing_map if key not in incoming_key_set]
            ids_to_remove = {rec.id for rec in to_delete}
            if ids_to_remove:
                # 先删掉「父在待删列表」的子行（同文档内），避免 FK 1451（含跨文档脏数据或未进 to_delete 的引用）
                r = db.execute(
                    delete(RFIGroundField).where(
                        RFIGroundField.document_number == doc,
                        RFIGroundField.parent_id.in_(ids_to_remove),
                    )
                )
                deleted_count += r.rowcount
            # 再按 level 从大到小删除待删行本身（部分可能已被上一步删掉，用 rowcount 避免重复计数）
            levels_to_delete = sorted({rec.level for rec in to_delete if rec.level is not None}, reverse=True)
            for lvl in levels_to_delete:
                ids_at_level = [rec.id for rec in to_delete if rec.level == lvl]
                if ids_at_level:
                    r2 = db.execute(delete(RFIGroundField).where(RFIGroundField.id.in_(ids_at_level)))
                    deleted_count += r2.rowcount

        db.commit()
        print("已更新/插入/删除 rfi_groundfields: 更新", updated_count, "条, 新增", inserted_count, "条, 删除", deleted_count, "条")
        if updated_details:
            dup_in_doc = sum(1 for _, _, _, r in updated_details if "同文档内重复编号" in r)
            match_db = len(updated_details) - dup_in_doc
            print("  更新说明（共 %d 条）：其中 同文档内重复编号 %d 条、与库内已有记录匹配 %d 条" % (len(updated_details), dup_in_doc, match_db))
            print("  明细（便于排查 Word 是否有多处相同编号或跨页重复）：")
            for doc, lvl, k, reason in updated_details:
                key_repr = f"level={lvl} itp_id/section={k[1]!r}"
                short = "重复编号" if "同文档内重复编号" in reason else "匹配库内"
                print(f"    - document_number={doc} {key_repr} [%s]" % short)
                print(f"      原因: {reason}")
    except Exception as e:
        db.rollback()
        raise
    finally:
        db.close()


def main():
    parser = argparse.ArgumentParser(description="从原版 ITP Word 文档导入 itp_definitions 和 rfi_groundfields")
    parser.add_argument("path", nargs="?", default=r"D:\Inspections\ITP\word", help="Word 文件所在目录")
    parser.add_argument("--preview", action="store_true", help="仅预览每个文档中 6/7 之间的表格")
    parser.add_argument("--dry-run", action="store_true", help="不提交数据库，仅打印将写入的数据")
    parser.add_argument("--verbose", "-v", action="store_true", help="打印每个文件的表格/ground 行数，便于排查")
    parser.add_argument("--no-skip-itp", action="store_true", help="已存在的 ITP 也更新")
    parser.add_argument("--file", type=str, default=None, help="只处理指定文件名（如 xxx.docx）")
    args = parser.parse_args()

    dir_path = Path(args.path)
    if not dir_path.is_dir():
        print(f"不是目录或不存在: {dir_path}")
        return
    if args.preview:
        preview_word_dir(dir_path, single_file=args.file)
        return
    import_from_word(
        dir_path,
        single_file=args.file,
        dry_run=args.dry_run,
        skip_existing_itp=not args.no_skip_itp,
        verbose=args.verbose,
    )


if __name__ == "__main__":
    main()
