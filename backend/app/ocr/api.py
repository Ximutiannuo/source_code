"""
OCR 工具 API - 图片文字识别接口

独立模块，不依赖业务权限，可作为通用工具使用。
"""
from fastapi import APIRouter, File, UploadFile, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Any, List, Literal
import os
import io

from app.ocr.service import (
    recognize,
    recognize_as_plain_text,
    recognize_structure,
    recognize_bordered_table,
    blocks_to_full_text,
)

router = APIRouter()

# 允许的输入格式：图片 + PDF
ALLOWED_CONTENT_TYPES = {
    "image/jpeg",
    "image/png",
    "image/bmp",
    "image/gif",
    "image/webp",
    "application/pdf",
}
ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".pdf"}

# PDF 最大 20MB，图片 5MB
MAX_IMAGE_SIZE = 5 * 1024 * 1024
MAX_PDF_SIZE = 20 * 1024 * 1024


def _is_pdf(contents: bytes, ct: str, ext: str) -> bool:
    if ct == "application/pdf" or ext == ".pdf":
        return len(contents) >= 4 and contents[:4] == b"%PDF"
    return False


def _max_file_size(contents: bytes, ct: str, ext: str) -> int:
    return MAX_PDF_SIZE if _is_pdf(contents, ct, ext) else MAX_IMAGE_SIZE


class OcrTextBlock(BaseModel):
    """单个识别文本块"""
    text: str
    confidence: float
    box: List[List[float]]


class OcrRecognizeResponse(BaseModel):
    """OCR 识别结果"""
    blocks: List[OcrTextBlock]
    full_text: str


@router.post("/recognize", response_model=OcrRecognizeResponse)
def ocr_recognize(
    file: UploadFile = File(..., description="图片文件 (jpg/png/bmp/gif/webp) 或 PDF"),
    lang: Literal["ch", "en", "ru", "en_ru", "ch_ru"] = "ch",
    format: Literal["blocks", "text"] = "blocks",
    model_type: Literal["server", "mobile"] = "mobile",
    noise_check: bool = True,
):
    """上传图片进行文字识别 (全 Mobile 引擎，兼顾速度与精度)

    - **lang**: ch=中英, en=英文, ru=俄文, en_ru=英俄, ch_ru=中俄
    - **format**: blocks=返回结构化块(含坐标), text=仅返回合并纯文本
    - **model_type**: 已弃用，后端统一使用 Mobile，保留参数仅兼容前端
    - **noise_check**: 是否启用噪点检查（移除箭头、格式符号等误识字符），默认 True
    """
    # 校验文件类型
    ct = (file.content_type or "").lower()
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ct not in ALLOWED_CONTENT_TYPES and ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的格式，允许: 图片 (jpg/png/bmp/gif/webp) 或 PDF",
        )

    contents = file.file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="文件内容为空")
    max_size = _max_file_size(contents, ct, ext)
    if len(contents) > max_size:
        raise HTTPException(status_code=400, detail=f"文件过大，最大支持 {max_size // 1024 // 1024}MB")

    try:
        if _is_pdf(contents, ct, ext):
            # PDF 与 OCR 分离：此处仅做「转图像 + OCR」，不做 PDF 元素提取
            from app.ocr.pdf_input import pdf_pages_to_images
            imgs = pdf_pages_to_images(contents)
            if not imgs:
                raise HTTPException(status_code=500, detail="PDF 无法转为图像，请检查文件或使用「PDF 提取」模式（仅电子版）")
            all_blocks = []
            y_offset = 0
            for img in imgs:
                page_blocks = recognize(img, lang=lang, model_type=model_type, enable_noise_filter=noise_check)
                for b in page_blocks:
                    nb = {"text": b["text"], "confidence": b["confidence"], "box": [[p[0], p[1] + y_offset] for p in (b.get("box") or [])]}
                    all_blocks.append(nb)
                y_offset += (img.shape[0] if hasattr(img, "shape") else 0)
            blocks = all_blocks
            full_text = blocks_to_full_text(blocks, lang=lang, enable_noise_filter=noise_check)
        else:
            blocks = recognize(
                contents, lang=lang, model_type=model_type, enable_noise_filter=noise_check
            )
            full_text = blocks_to_full_text(blocks, lang=lang, enable_noise_filter=noise_check)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR 识别失败: {str(e)}")

    if format == "text":
        return OcrRecognizeResponse(blocks=[], full_text=full_text)

    ocr_blocks = [
        OcrTextBlock(text=b["text"], confidence=b["confidence"], box=b.get("box", []))
        for b in blocks
    ]
    return OcrRecognizeResponse(blocks=ocr_blocks, full_text=full_text)


class OcrStructureResponse(BaseModel):
    """PP-StructureV3 文档解析结果（含表格结构）"""
    markdown: str
    tables: List[List[List[str]]]


class OcrBorderedTableResponse(BaseModel):
    """带边框表格识别结果：网格检测 + OCR 单元格分配，格式完美匹配"""
    markdown: str
    tables: List[List[List[str]]]
    merges: List[List[int]]
    grid_detected: bool


class OcrPdfExtractResponse(BaseModel):
    """PDF 元素提取结果（仅电子版 PDF，不做 OCR）"""
    markdown: str
    tables: List[List[List[str]]]
    page_count: int
    scanned_or_empty: bool  # True 表示提取内容很少，可能是扫描版，建议改用 OCR 模式


@router.post("/recognize/pdf-extract", response_model=OcrPdfExtractResponse)
def ocr_recognize_pdf_extract(
    file: UploadFile = File(..., description="PDF 文件（仅电子版，扫描版请用 OCR 模式）"),
):
    """
    PDF 元素提取：仅从 PDF 抓取文本与表格，不做 OCR。
    适用于电子版 PDF。若返回 scanned_or_empty=True，表示可能是扫描版，请改用「通用文字识别」或「带边框表格」并上传同一 PDF（将自动转图像后 OCR）。
    """
    ct = (file.content_type or "").lower()
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ct != "application/pdf" and ext != ".pdf":
        raise HTTPException(status_code=400, detail="仅支持 PDF 文件")
    contents = file.file.read()
    if not contents or len(contents) < 4 or contents[:4] != b"%PDF":
        raise HTTPException(status_code=400, detail="无效的 PDF 文件")
    if len(contents) > MAX_PDF_SIZE:
        raise HTTPException(status_code=400, detail=f"PDF 过大，最大支持 {MAX_PDF_SIZE // 1024 // 1024}MB")
    from app.ocr.pdf_input import extract_from_pdf
    data = extract_from_pdf(contents)
    if not data:
        return OcrPdfExtractResponse(
            markdown="",
            tables=[],
            page_count=0,
            scanned_or_empty=True,
        )
    full_text = (data.get("full_text") or "").strip()
    tables = data.get("tables") or []
    page_count = data.get("page_count") or 0
    scanned_or_empty = len(full_text) < 80 and not tables
    return OcrPdfExtractResponse(
        markdown=full_text,
        tables=tables,
        page_count=page_count,
        scanned_or_empty=scanned_or_empty,
    )


@router.post("/recognize/structure", response_model=OcrStructureResponse)
def ocr_recognize_structure(
    file: UploadFile = File(..., description="图片或 PDF（文档/表格）"),
):
    """PP-StructureV3 文档解析；PDF 时将转成图像后解析（扫描版 PDF 也适用）。"""
    ct = (file.content_type or "").lower()
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ct not in ALLOWED_CONTENT_TYPES and ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="不支持的格式，允许: 图片或 PDF")
    contents = file.file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="文件内容为空")
    if len(contents) > _max_file_size(contents, ct, ext):
        raise HTTPException(status_code=400, detail="文件过大")
    try:
        if _is_pdf(contents, ct, ext):
            # PDF 与 OCR 分离：此处仅做「转图像 + 结构解析」，不做 PDF 元素提取
            from app.ocr.pdf_input import pdf_pages_to_images
            imgs = pdf_pages_to_images(contents, max_pages=1)
            if not imgs:
                raise HTTPException(status_code=500, detail="PDF 无法转为图像，请使用「PDF 提取」模式（仅电子版）或检查文件")
            data = recognize_structure(imgs[0])
        else:
            data = recognize_structure(contents)
    except HTTPException:
        raise
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")
    return OcrStructureResponse(markdown=data["markdown"], tables=data["tables"])


@router.post("/recognize/bordered-table", response_model=OcrBorderedTableResponse)
def ocr_recognize_bordered_table(
    file: UploadFile = File(..., description="带边框的表格图片或 PDF（如俄语/英语技术规格表）"),
    lang: Literal["ch", "en", "ru", "en_ru", "ch_ru"] = "ru",
):
    """
    带边框表格识别：利用形态学检测横竖线得到网格，再 OCR 并将文本分配到单元格，
    实现从识别到表格式的完美匹配。支持合并单元格检测（merges）。
    若网格检测失败则自动回退为按行分组的表格结果。
    """
    ct = (file.content_type or "").lower()
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ct not in ALLOWED_CONTENT_TYPES and ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="不支持的格式，允许: 图片或 PDF")
    contents = file.file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="文件内容为空")
    if len(contents) > _max_file_size(contents, ct, ext):
        raise HTTPException(status_code=400, detail="文件过大")
    try:
        if _is_pdf(contents, ct, ext):
            from app.ocr.pdf_input import pdf_pages_to_images
            imgs = pdf_pages_to_images(contents, max_pages=1)
            if not imgs:
                raise HTTPException(status_code=500, detail="PDF 无法转为图像")
            contents = imgs[0]
        data = recognize_bordered_table(contents, lang=lang)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"带边框表格识别失败: {str(e)}")
    return OcrBorderedTableResponse(
        markdown=data["markdown"],
        tables=data["tables"],
        merges=[list(m) for m in data["merges"]],
        grid_detected=data["grid_detected"],
    )


class OcrExportBorderedTableXlsxRequest(BaseModel):
    """导出带边框表格为 XLSX 的请求体"""
    tables: List[List[List[str]]]
    merges: List[List[int]] = []  # 每项为 [r0, c0, r1, c1]（0-based，含首含尾）
    flatten: bool = False  # True 时不合并单元格，输出扁平表格，避免合并导致内容丢失


class OcrExportDocxRequest(BaseModel):
    """导出 OCR 结果为 Word 的请求体"""
    markdown: str = ""
    tables: List[List[List[str]]] = []


# Excel XML 不允许的控制字符（会导致「文件已损坏」）
_ILLEGAL_XML_CHARS_RE = __import__("re").compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
# 欧式小数（逗号作小数点，如 13,658 = 13.658），需强制文本格式保留原样
_EUROPEAN_DECIMAL_RE = __import__("re").compile(r"^\d+,\d+$")


def _sanitize_cell_value(val: Any) -> str:
    """移除非法字符，并将单换行转为空格实现自然换行（段落间保留 \\n）。"""
    if val is None:
        return ""
    s = str(val)
    # 移除 Excel XML 非法控制字符
    s = _ILLEGAL_XML_CHARS_RE.sub("", s)
    # 自然换行：单换行视为行内换行→空格，双换行视为段落→保留一个 \\n
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    parts = s.split("\n\n")
    parts = [" ".join(p.split("\n")) for p in parts]
    return "\n".join(parts)


def _consolidate_merge_regions(
    table: List[List[str]], merges: List[List[int]]
) -> List[List[str]]:
    """
    在写入 Excel 前，将每个合并区域内的所有单元格文本汇总到左上角，
    避免 merge_cells 后只保留左上角导致内容丢失。
    """
    if not table or not merges:
        return table
    # 深拷贝，避免修改原始数据
    out = [list(row) for row in table]
    for m in merges:
        if len(m) != 4:
            continue
        r0, c0, r1, c1 = m[0], m[1], m[2], m[3]
        if r1 < r0 or c1 < c0:
            continue
        if r0 == r1 and c0 == c1:
            continue
        parts = []
        for r in range(r0, r1 + 1):
            for c in range(c0, c1 + 1):
                if r < len(out) and c < len(out[r]):
                    v = (out[r][c] or "").strip()
                    if v:
                        parts.append(v)
        # 汇总到左上角：非空部分用换行拼接，若都为空则保留原左上角
        if parts:
            out[r0][c0] = "\n".join(parts)
        for r in range(r0, r1 + 1):
            for c in range(c0, c1 + 1):
                if (r, c) != (r0, c0) and r < len(out) and c < len(out[r]):
                    out[r][c] = ""
    return out


def _build_bordered_table_xlsx(
    tables: List[List[List[str]]],
    merges: List[List[int]],
    flatten: bool = False,
) -> bytes:
    """
    生成带框线、合并单元格、统一字体的 XLSX 字节流。
    支持多表（分页导致）：每表一个工作表（表格1、表格2...），merges 仅作用于第一个表。
    """
    import openpyxl
    from openpyxl.styles import Font, Border, Side, Alignment

    if not tables:
        wb = openpyxl.Workbook()
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.getvalue()

    wb = openpyxl.Workbook()
    thin = Side(style="thin", color="000000")
    cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)
    default_font = Font(name="Arial", size=11)
    header_font = Font(name="Arial", size=11, bold=True)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for idx, tbl in enumerate(tables):
        if not tbl:
            continue
        use_merges = [] if flatten else (merges if idx == 0 else [])
        table = _consolidate_merge_regions(tbl, use_merges)
        row_count = len(table)
        col_count = max(len(row) for row in table) if table else 0
        if col_count == 0:
            continue

        ws_name = f"表格{idx + 1}" if len(tables) > 1 else "表格"
        if idx == 0:
            ws = wb.active
            ws.title = ws_name
        else:
            ws = wb.create_sheet(title=ws_name)

        for r, row in enumerate(table, start=1):
            for c in range(col_count):
                cell = ws.cell(row=r, column=c + 1)
                raw = row[c] if c < len(row) else ""
                val = _sanitize_cell_value(raw)
                cell.value = val
                if val and _EUROPEAN_DECIMAL_RE.match(val.strip()):
                    cell.number_format = "@"  # 强制文本，保留 13,658 原样（逗号作小数点）
                cell.border = cell_border
                cell.font = header_font if r == 1 else default_font
                cell.alignment = center_align

        for m in use_merges:
            if len(m) != 4:
                continue
            r0, c0, r1, c1 = m[0], m[1], m[2], m[3]
            if r1 < r0 or c1 < c0 or (r0 == r1 and c0 == c1):
                continue
            ws.merge_cells(
                start_row=r0 + 1,
                start_column=c0 + 1,
                end_row=r1 + 1,
                end_column=c1 + 1,
            )
            ws.cell(row=r0 + 1, column=c0 + 1).alignment = center_align

        for c in range(1, col_count + 1):
            try:
                col_letter = openpyxl.utils.get_column_letter(c)
                w = min(50, max(12, len(str(table[0][c - 1])) + 2)) if table and c <= len(table[0]) else 14
                ws.column_dimensions[col_letter].width = w
            except Exception:
                pass

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_docx(markdown: str, tables: List[List[List[str]]]) -> bytes:
    """生成 Word 文档：正文（markdown 纯文本） + 所有表格。"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 正文
    if markdown and markdown.strip():
        for para in markdown.strip().split("\n\n"):
            p = doc.add_paragraph(para.strip().replace("\n", " "))
            p.paragraph_format.space_after = Pt(6)

    # 表格（分页导致的多表全部导出）
    for idx, tbl in enumerate(tables):
        if not tbl:
            continue
        if idx > 0 or (markdown and markdown.strip()):
            doc.add_paragraph()
        rows, cols = len(tbl), max(len(r) for r in tbl) if tbl else 0
        if cols == 0:
            continue
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for r, row in enumerate(tbl):
            for c in range(cols):
                cell = table.rows[r].cells[c]
                cell.text = _sanitize_cell_value(row[c] if c < len(row) else "")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/export/docx")
def ocr_export_docx(body: OcrExportDocxRequest) -> Response:
    """将 OCR 识别结果导出为 Word（正文 + 所有表格）。"""
    try:
        docx_bytes = _build_docx(body.markdown, body.tables)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成 Word 失败: {str(e)}")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=ocr_result.docx"},
    )


@router.post("/export/bordered-table-xlsx")
def ocr_export_bordered_table_xlsx(body: OcrExportBorderedTableXlsxRequest) -> Response:
    """
    将带边框表格识别结果导出为 XLSX。
    支持合并单元格、框线、统一字体与对齐，使导出效果接近金山 PDF 等工具的表格导出。
    """
    try:
        xlsx_bytes = _build_bordered_table_xlsx(
            body.tables, body.merges, flatten=body.flatten
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成 XLSX 失败: {str(e)}")
    return Response(
        content=xlsx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": "attachment; filename=ocr_bordered_table.xlsx",
        },
    )


@router.post("/recognize/text")
def ocr_recognize_text_only(
    file: UploadFile = File(..., description="图片文件或 PDF"),
    lang: Literal["ch", "en", "ru", "en_ru", "ch_ru"] = "ch",
) -> dict:
    """仅返回合并后的纯文本；PDF 时优先抓取元素提取文本，失败再 OCR。"""
    ct = (file.content_type or "").lower()
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ct not in ALLOWED_CONTENT_TYPES and ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="不支持的格式，允许: 图片或 PDF")
    contents = file.file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="文件内容为空")
    if len(contents) > _max_file_size(contents, ct, ext):
        raise HTTPException(status_code=400, detail="文件过大")
    try:
        if _is_pdf(contents, ct, ext):
            # PDF 与 OCR 分离：此处仅做「转图像 + OCR」
            from app.ocr.pdf_input import pdf_pages_to_images
            imgs = pdf_pages_to_images(contents)
            if not imgs:
                raise HTTPException(status_code=500, detail="PDF 无法转为图像，请使用「PDF 提取」模式（仅电子版）或检查文件")
            all_blocks = []
            y_offset = 0
            for img in imgs:
                page_blocks = recognize(img, lang=lang, enable_noise_filter=True)
                for b in page_blocks:
                    nb = {"text": b["text"], "confidence": b["confidence"], "box": [[p[0], p[1] + y_offset] for p in (b.get("box") or [])]}
                    all_blocks.append(nb)
                y_offset += (img.shape[0] if hasattr(img, "shape") else 0)
            text = blocks_to_full_text(all_blocks, lang=lang, enable_noise_filter=True)
            return {"text": text}
        text = recognize_as_plain_text(contents, lang=lang)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR 识别失败: {str(e)}")
    return {"text": text}
